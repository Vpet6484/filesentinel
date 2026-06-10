from docx import Document
from docx.shared import Pt, RGBColor
import tempfile


# ==================================================
# DOCX HIDDEN TEXT WATERMARK (100% RELIABLE)
# ==================================================

def embed_docx(src_path, token):

    doc = Document(src_path)

    # Create hidden paragraph
    p = doc.add_paragraph()

    run = p.add_run(f"FS-WM:{token}")

    # Make invisible
    run.font.size = Pt(1)
    run.font.color.rgb = RGBColor(255, 255, 255)
    run.font.hidden = True


    temp_path = tempfile.mktemp(suffix=".docx")

    doc.save(temp_path)

    return temp_path
